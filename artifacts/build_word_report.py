from pathlib import Path
import re

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORT_MD = ROOT / "PROJECT_ANALYSIS_AND_TEST_REPORT.md"
OUT_DOCX = ROOT / "CJLU_Student_App_Full_Project_Report.docx"
DIAGRAM_DIR = ROOT / "artifacts" / "diagrams"
SCREEN_DIR = ROOT / "artifacts" / "screenshots"

NAVY = "0B2545"
BLUE = "0F6FFF"
TEAL = "18A999"
PALE_BLUE = "E8F0FF"
PALE_TEAL = "E3FBF7"
LIGHT = "F4F6F9"
MID = "D7E0EA"
DARK = "18202A"
MUTED = "5A6675"
RED = "9B1C1C"


def font(size, bold=False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "",
        "/Library/Fonts/Arial.ttf",
    ]
    if bold:
        candidates = [
            "/System/Library/Fonts/Supplemental/Arial Bold.ttf",
            "/System/Library/Fonts/Supplemental/Arial.ttf",
        ]
    for candidate in candidates:
        if candidate and Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


F_TITLE = font(42, True)
F_H = font(29, True)
F_BODY = font(23)
F_SMALL = font(19)


def pil_color(value):
    if isinstance(value, str) and re.fullmatch(r"[0-9A-Fa-f]{6}", value):
        return "#" + value
    return value


def rounded_box(draw, xy, title, lines, fill, outline=BLUE, title_fill=None):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(
        xy,
        radius=24,
        fill=pil_color(fill),
        outline=pil_color(outline),
        width=4,
    )
    if title_fill:
        draw.rounded_rectangle(
            (x1, y1, x2, y1 + 58),
            radius=22,
            fill=pil_color(title_fill),
        )
        draw.rectangle((x1, y1 + 32, x2, y1 + 58), fill=pil_color(title_fill))
        title_color = "white"
    else:
        title_color = pil_color(NAVY)
    draw.text((x1 + 22, y1 + 15), title, font=F_H, fill=title_color)
    y = y1 + 76
    for line in lines:
        draw.text((x1 + 22, y), line, font=F_SMALL, fill=pil_color(DARK))
        y += 30


def arrow(draw, start, end, color=BLUE, width=6, label=None, label_pos=None):
    color = pil_color(color)
    draw.line((start, end), fill=color, width=width)
    x1, y1 = start
    x2, y2 = end
    import math
    angle = math.atan2(y2 - y1, x2 - x1)
    length = 18
    for offset in (2.55, -2.55):
        px = x2 + length * math.cos(angle + offset)
        py = y2 + length * math.sin(angle + offset)
        draw.line((x2, y2, px, py), fill=color, width=width)
    if label:
        lx, ly = label_pos or ((x1 + x2) // 2, (y1 + y2) // 2)
        bbox = draw.textbbox((lx, ly), label, font=F_SMALL)
        draw.rounded_rectangle(
            (bbox[0] - 8, bbox[1] - 5, bbox[2] + 8, bbox[3] + 5),
            radius=8,
            fill="white",
        )
        draw.text((lx, ly), label, font=F_SMALL, fill=pil_color(MUTED))


def canvas(title, subtitle):
    img = Image.new("RGB", (1600, 920), "white")
    draw = ImageDraw.Draw(img)
    draw.rectangle((0, 0, 1600, 110), fill=pil_color(NAVY))
    draw.text((55, 24), title, font=F_TITLE, fill="white")
    draw.text((58, 121), subtitle, font=F_BODY, fill=pil_color(MUTED))
    return img, draw


def make_diagrams():
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)

    img, d = canvas(
        "CJLU Student App - System Architecture",
        "Three-module Gradle system with mobile, shared contract, and server-side administration",
    )
    rounded_box(d, (70, 220, 500, 480), "Android Client", [
        "Jetpack Compose + Material 3",
        "Room local database",
        "Retrofit / OkHttp REST",
        "Ktor WebSocket client",
        "FCM + Glance widget",
    ], PALE_BLUE, BLUE, BLUE)
    rounded_box(d, (585, 260, 1015, 460), "Shared Contract", [
        "kotlinx.serialization DTOs",
        "Authentication and profiles",
        "Requests and messages",
        "Academic data and push events",
    ], LIGHT, NAVY, NAVY)
    rounded_box(d, (1100, 220, 1530, 480), "Ktor Backend", [
        "REST and WebSocket routes",
        "JWT, API key, BCrypt",
        "Exposed ORM",
        "H2 / PostgreSQL",
        "FreeMarker admin portal",
    ], PALE_TEAL, TEAL, TEAL)
    rounded_box(d, (180, 620, 620, 810), "Student Experience", [
        "Login, dashboard, academics",
        "Services, uploads, messages",
        "Offline cached rendering",
    ], "FFFFFF", MID)
    rounded_box(d, (980, 620, 1420, 810), "Staff Experience", [
        "Request administration",
        "Academic corrections",
        "Targeted announcements",
    ], "FFFFFF", MID)
    arrow(d, (500, 340), (585, 340), label="uses", label_pos=(515, 300))
    arrow(d, (1015, 340), (1100, 340), label="uses", label_pos=(1025, 300))
    arrow(d, (500, 505), (1100, 505), color=TEAL, label="REST / WebSocket", label_pos=(690, 470))
    arrow(d, (400, 620), (330, 480), color=NAVY)
    arrow(d, (1200, 480), (1200, 620), color=NAVY)
    img.save(DIAGRAM_DIR / "architecture.png")

    img, d = canvas(
        "Application Data Flow",
        "User action, state coordination, local persistence, network request, and UI update",
    )
    nodes = [
        ((55, 300, 320, 500), "Compose Screen", ["User event", "State rendering"], PALE_BLUE, BLUE),
        ((390, 300, 670, 500), "Coordinator", ["MainActivity", "Navigation callbacks"], LIGHT, NAVY),
        ((740, 210, 1020, 410), "Repository", ["Request / Message", "Academic / Catalog"], PALE_TEAL, TEAL),
        ((740, 520, 1020, 720), "Room Cache", ["Flows to UI", "Offline fallback"], "FFF8E8", "B7791F"),
        ((1090, 300, 1370, 500), "Retrofit API", ["JWT + API key", "JSON / multipart"], PALE_BLUE, BLUE),
        ((1410, 300, 1580, 500), "Ktor", ["Routes", "Database"], PALE_TEAL, TEAL),
    ]
    for xy, title, lines, fill, outline in nodes:
        rounded_box(d, xy, title, lines, fill, outline)
    arrow(d, (320, 400), (390, 400), label="event", label_pos=(330, 360))
    arrow(d, (670, 370), (740, 310), label="load", label_pos=(665, 285))
    arrow(d, (880, 410), (880, 520), color="B7791F", label="save", label_pos=(900, 450))
    arrow(d, (740, 620), (600, 500), color="B7791F", label="Flow", label_pos=(610, 545))
    arrow(d, (1020, 310), (1090, 370), label="fetch", label_pos=(1020, 275))
    arrow(d, (1370, 400), (1410, 400))
    arrow(d, (1410, 455), (1020, 455), color=TEAL, label="response", label_pos=(1170, 425))
    arrow(d, (390, 455), (320, 455), color=NAVY, label="state", label_pos=(320, 500))
    img.save(DIAGRAM_DIR / "data-flow.png")

    img, d = canvas(
        "Realtime Synchronization Flow",
        "The same logical update can arrive through WebSocket while active or FCM in the background",
    )
    rounded_box(d, (50, 310, 410, 540), "Admin / Server Event", [
        "Message sent",
        "Request status changed",
        "Academic record updated",
    ], PALE_TEAL, TEAL, TEAL)
    rounded_box(d, (470, 205, 790, 405), "WebSocket Hub", [
        "Connected student session",
        "Typed JSON push payload",
    ], PALE_BLUE, BLUE)
    rounded_box(d, (470, 560, 790, 760), "Firebase Cloud", [
        "Data message mirrors push",
        "Used when app is backgrounded",
    ], "FFF8E8", "B7791F")
    rounded_box(d, (900, 310, 1210, 540), "Android Handlers", [
        "RealtimePushHandler",
        "FCM push handler",
        "RealtimeSyncCoordinator",
    ], LIGHT, NAVY)
    rounded_box(d, (1310, 310, 1550, 540), "Effects", [
        "Refresh Room data",
        "Invalidate cache",
        "Update widget",
        "Show notification",
    ], PALE_TEAL, TEAL)
    arrow(d, (410, 380), (470, 310), label="foreground", label_pos=(375, 300))
    arrow(d, (410, 480), (470, 650), color="B7791F", label="background", label_pos=(360, 570))
    arrow(d, (790, 310), (900, 380))
    arrow(d, (790, 660), (900, 470), color="B7791F")
    arrow(d, (1210, 425), (1310, 425), color=TEAL)
    img.save(DIAGRAM_DIR / "realtime-sync.png")

    img, d = canvas(
        "Academic Cache Invalidation Flow",
        "Targeted invalidation prevents stale academic screens while preserving offline fallback",
    )
    rounded_box(d, (55, 315, 335, 535), "Push Payload", [
        "academic_updated",
        "scope identifies data",
    ], PALE_BLUE, BLUE)
    rounded_box(d, (405, 315, 720, 535), "Sync Coordinator", [
        "Parse action",
        "Select affected keys",
    ], LIGHT, NAVY)
    rounded_box(d, (790, 210, 1110, 420), "Room Invalidation", [
        "attendance",
        "timetable",
        "transcript",
        "dormitory / calendar",
    ], "FFF8E8", "B7791F")
    rounded_box(d, (790, 540, 1110, 750), "UI Refresh Nonce", [
        "Triggers screen reload",
        "Profile state refreshed",
    ], PALE_TEAL, TEAL)
    rounded_box(d, (1190, 315, 1545, 535), "Repository Reload", [
        "Try API first",
        "Write fresh JSON to Room",
        "Use remaining cache on failure",
    ], PALE_BLUE, BLUE)
    arrow(d, (335, 425), (405, 425))
    arrow(d, (720, 380), (790, 315), color="B7791F")
    arrow(d, (720, 470), (790, 645), color=TEAL)
    arrow(d, (1110, 315), (1190, 380))
    arrow(d, (1110, 645), (1190, 475), color=TEAL)
    d.text(
        (245, 825),
        "Outcome: affected records are refreshed; unrelated cached data remains available.",
        font=F_BODY,
        fill=pil_color(NAVY),
    )
    img.save(DIAGRAM_DIR / "cache-invalidation.png")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    total_dxa = sum(int(width * 1440) for width in widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(grid_col)
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(row.cells[idx])


def set_run(run, size=None, bold=None, color=None, italic=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    specs = {
        "Title": (26, NAVY, 0, 8),
        "Subtitle": (13, MUTED, 0, 8),
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, NAVY, 8, 4),
    }
    for name, (size, color, before, after) in specs.items():
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.10


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


def configure_page(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run("CJLU STUDENT APP  |  PROJECT REPORT")
    set_run(r, size=8.5, bold=True, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("Page ")
    set_run(r, size=9, color=MUTED)
    add_page_number(footer)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=9, italic=True, color=MUTED)


def add_figure(doc, image, caption, width=6.35):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    shape = p.add_run().add_picture(str(image), width=Inches(width))
    shape._inline.docPr.set("descr", caption)
    shape._inline.docPr.set("title", caption)
    add_caption(doc, caption)


def add_callout(doc, title, text, fill=PALE_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title + ": ")
    set_run(r, bold=True, color=NAVY)
    r = p.add_run(text)
    set_run(r, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_markdown_table(doc, rows):
    parsed = [[cell.strip() for cell in row.strip().strip("|").split("|")] for row in rows]
    if len(parsed) > 1 and all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in parsed[1]):
        parsed.pop(1)
    if not parsed:
        return
    cols = max(len(row) for row in parsed)
    table = doc.add_table(rows=len(parsed), cols=cols)
    widths = [6.5 / cols] * cols
    set_table_widths(table, widths)
    mark_header_row(table.rows[0])
    for r_idx, values in enumerate(parsed):
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if r_idx == 0:
                set_cell_shading(cell, LIGHT)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            text = values[c_idx] if c_idx < len(values) else ""
            add_inline_markdown(p, text)
            if r_idx == 0:
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_markdown(doc, markdown):
    lines = markdown.splitlines()
    in_code = False
    code_lines = []
    idx = 0
    while idx < len(lines):
        line = lines[idx]
        if line.startswith("```"):
            if in_code:
                table = doc.add_table(rows=1, cols=1)
                set_table_widths(table, [6.5])
                set_cell_shading(table.cell(0, 0), "F6F8FA")
                p = table.cell(0, 0).paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                r = p.add_run("\n".join(code_lines))
                r.font.name = "Courier New"
                r.font.size = Pt(8.5)
                code_lines = []
                in_code = False
            else:
                in_code = True
            idx += 1
            continue
        if in_code:
            code_lines.append(line)
            idx += 1
            continue
        if not line.strip() or line.startswith("# CJLU Student App:"):
            idx += 1
            continue
        if line.startswith("|"):
            table_rows = []
            while idx < len(lines) and lines[idx].startswith("|"):
                table_rows.append(lines[idx])
                idx += 1
            add_markdown_table(doc, table_rows)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=3)
        elif re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            text = re.sub(r"^\d+\. ", "", line)
            add_inline_markdown(p, text)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_markdown(p, line[2:])
        elif line.startswith("```mermaid"):
            idx += 1
            continue
        elif line.startswith("**") and line.endswith("**") and ":" in line:
            p = doc.add_paragraph()
            add_inline_markdown(p, line)
        else:
            p = doc.add_paragraph()
            add_inline_markdown(p, line)
        idx += 1


def add_inline_markdown(paragraph, text):
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            set_run(r, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1])
            r.font.name = "Courier New"
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor.from_string(NAVY)
        else:
            r = paragraph.add_run(part)
            set_run(r)


def add_demo_screens(doc):
    doc.add_heading("Application Demonstration Screenshots", level=1)
    p = doc.add_paragraph(
        "The following images were captured from the real debug APK running on a Pixel 4a Android emulator. "
        "The offline banners demonstrate the implemented fallback state when the configured remote service is unavailable."
    )
    p.paragraph_format.space_after = Pt(10)
    shots = [
        ("01-login.png", "Student login and credential guidance"),
        ("02-home.png", "Home dashboard with attendance and quick actions"),
        ("03-services.png", "Service Center with search, categories, and popular services"),
        ("04-messages.png", "Campus Messages empty/offline state"),
        ("05-profile.png", "Student profile, identity data, and notification settings"),
        ("07-service-detail.png", "Visa Extension request history and new-request workflow"),
    ]
    for idx in range(0, len(shots), 2):
        table = doc.add_table(rows=2, cols=2)
        set_table_widths(table, [3.25, 3.25])
        for col, item in enumerate(shots[idx:idx + 2]):
            name, caption = item
            image_cell = table.cell(0, col)
            image_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = image_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            shape = p.add_run().add_picture(str(SCREEN_DIR / name), width=Inches(2.62))
            shape._inline.docPr.set("descr", caption)
            shape._inline.docPr.set("title", caption)
            cap_cell = table.cell(1, col)
            cap_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cap_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(caption)
            set_run(r, size=9, italic=True, color=MUTED)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_admin_portal_demo(doc):
    doc.add_page_break()
    doc.add_heading("Backend Administration Portal Demonstration", level=1)
    p = doc.add_paragraph(
        "The backend includes a browser-based staff console served directly by the Ktor application with "
        "FreeMarker templates. The following screenshots were captured from the running local backend at "
        "`http://localhost:8080/admin`. They demonstrate authenticated access, operational monitoring, and "
        "the administration-to-mobile realtime messaging workflow."
    )
    p.paragraph_format.space_after = Pt(8)

    add_callout(
        doc,
        "Demonstration scope",
        "Staff authenticate through a protected session, review service workload, manage student requests "
        "and academic data, and send broadcast or targeted inbox messages. Administrative changes are "
        "persisted by the backend and relevant student devices are prompted to refresh through WebSocket "
        "and Firebase Cloud Messaging channels.",
        PALE_TEAL,
    )

    add_figure(
        doc,
        SCREEN_DIR / "08-admin-login.png",
        "Figure 5. Backend administration login screen protected by form authentication and an admin session.",
        width=6.25,
    )
    add_figure(
        doc,
        SCREEN_DIR / "09-admin-dashboard.png",
        "Figure 6. Administration overview showing request status totals, service-catalog coverage, and category workload.",
        width=6.25,
    )

    doc.add_heading("Administration Workflow Demonstrated", level=2)
    workflow = [
        ("1. Authenticate", "Staff sign in and receive the protected ADMIN_SESSION cookie."),
        ("2. Monitor", "The overview summarizes total, submitted, in-review, action-needed, and completed requests."),
        ("3. Operate", "Dedicated panels manage requests, walk-in registration, attendance, transcripts, learning data, and the academic calendar."),
        ("4. Notify", "The inbox form sends a broadcast or student-targeted bilingual message and can mark it as requiring action."),
        ("5. Synchronize", "The backend stores the change, emits the relevant WebSocket event, and mirrors background delivery through FCM when configured."),
    ]
    table = doc.add_table(rows=1, cols=2)
    set_table_widths(table, [1.45, 5.05])
    table.rows[0].cells[0].text = "Stage"
    table.rows[0].cells[1].text = "Observed behavior"
    mark_header_row(table.rows[0])
    for cell in table.rows[0].cells:
        set_cell_shading(cell, LIGHT)
        for run in cell.paragraphs[0].runs:
            set_run(run, bold=True, color=NAVY)
    for stage, behavior in workflow:
        cells = table.add_row().cells
        cells[0].text = stage
        cells[1].text = behavior
        for cell in cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cells[0].paragraphs[0].runs:
            set_run(run, bold=True, color=BLUE)
        for run in cells[1].paragraphs[0].runs:
            set_run(run, color=DARK)

    doc.add_page_break()
    doc.add_heading("Realtime Student Inbox Administration", level=2)
    p = doc.add_paragraph(
        "This panel demonstrates the backend side of the mobile Messages feature. An administrator can select "
        "broadcast delivery or one student, choose a category, enter English content with optional Chinese "
        "translations, link the message to a service, and request an action indicator. Submitting the form "
        "writes the message to backend storage and notifies connected apps to refresh their inbox."
    )
    p.paragraph_format.space_after = Pt(8)
    add_figure(
        doc,
        SCREEN_DIR / "10-admin-inbox.png",
        "Figure 7. Student inbox administration form for targeted or broadcast bilingual messages and realtime app notification.",
        width=3.15,
    )


def build_docx():
    make_diagrams()
    doc = Document()
    configure_page(doc)
    configure_styles(doc)

    # Editorial cover pattern using the standard_business_brief preset.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(110)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("FULL PROJECT ANALYSIS")
    set_run(r, size=11, bold=True, color=TEAL)
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CJLU Student App")
    set_run(r, size=30, bold=True, color=NAVY)
    p = doc.add_paragraph(style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Architecture, Data Flow, Realtime Synchronization, Cache Invalidation, Testing, and UI Demonstration")
    set_run(r, size=14, color=MUTED)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(55)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Android Client  |  Shared Contract  |  Ktor Backend")
    set_run(r, size=11, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Prepared June 15, 2026")
    set_run(r, size=10.5, color=MUTED)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(90)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Source-code analysis, automated verification, emulator demonstration, and technical recommendations")
    set_run(r, size=10, italic=True, color=MUTED)
    doc.add_page_break()

    doc.add_heading("Report Guide", level=1)
    for item in [
        "Introduction",
        "Project Requirements",
        "System Architecture",
        "Application Data Flow",
        "Realtime Synchronization",
        "Cache Invalidation",
        "UI Design",
        "Model Class Design",
        "Development Tools",
        "Android App Demonstration",
        "Backend Administration Portal Demonstration",
        "Testing and Results",
        "Problems and Solutions",
        "Project Experience",
        "Summary and Conclusion",
    ]:
        doc.add_paragraph(item, style="List Number")
    add_callout(
        doc,
        "Verification status",
        "28 JVM tests passed, debug APK assembly passed, Android lint completed with 0 errors, and the Ktor backend started successfully. Device instrumentation tests were not executed as a complete suite.",
        PALE_TEAL,
    )

    doc.add_heading("Architecture and Operational Flows", level=1)
    add_figure(doc, DIAGRAM_DIR / "architecture.png", "Figure 1. Overall CJLU Student App system architecture.")
    add_figure(doc, DIAGRAM_DIR / "data-flow.png", "Figure 2. Application data flow from user interaction to local and remote data sources.")
    add_figure(doc, DIAGRAM_DIR / "realtime-sync.png", "Figure 3. Realtime synchronization through WebSocket and Firebase Cloud Messaging.")
    add_figure(doc, DIAGRAM_DIR / "cache-invalidation.png", "Figure 4. Targeted academic-cache invalidation and reload sequence.")

    add_demo_screens(doc)
    add_admin_portal_demo(doc)
    doc.add_page_break()

    markdown = REPORT_MD.read_text(encoding="utf-8")
    # Remove the original metadata and Mermaid block because the Word version has a richer cover and rendered drawings.
    markdown = re.sub(r"\*\*Project:\*\*.*?\n\n", "", markdown, flags=re.S)
    markdown = re.sub(r"```mermaid.*?```", "", markdown, flags=re.S)
    add_markdown(doc, markdown)

    doc.core_properties.title = "CJLU Student App - Full Project Analysis"
    doc.core_properties.subject = "Architecture, testing, data flow, realtime sync, cache invalidation, and UI demonstration"
    doc.core_properties.author = "CJLU Student App Project Team"
    doc.core_properties.keywords = "Android, Jetpack Compose, Ktor, Room, architecture, testing"
    doc.save(OUT_DOCX)
    return OUT_DOCX


if __name__ == "__main__":
    print(build_docx())
